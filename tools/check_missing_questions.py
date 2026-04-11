#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
原本Wordファイルで問49〜61の存在を確認

目的:
    - R1〜R7の各年度で問49〜61の解説が存在するか
    - R6年度の問題網羅性を確認
    - 法規問題（問49〜61）の扱いを理解する
"""

import sys
import io
from pathlib import Path
from docx import Document
import re

# UTF-8出力設定
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

ORIGINAL_FOLDER = Path(r"C:\Users\masam\Downloads\１級土木-20260214T233348Z-1-001\１級土木")

# 推奨ファイル（前セッション調査より）
FILES = {
    'R1': "1級土木 解説文集 令和元年度.docx",
    'R2': "1級土木 解説文集 令和2年度.docx",
    'R3': "1級土木 解説文集 令和3年度(1).docx",
    'R4': "1級土木 解説文集 令和4年度.docx",
    'R5': "1級土木 解説文集 令和5年度.docx",
    'R6': "1級土木 解説文集 令和6年度.docx",
    'R7': "1級土木 解説文集 令和7年度.docx",
}

def check_question_numbers(year, file_path):
    """
    Wordファイルから問題番号を抽出し、49〜61の存在を確認
    """
    doc = Document(file_path)

    # 問題番号パターン（複数形式に対応）
    patterns = [
        r'問(\d+)[\s（]',              # 問1（
        r'問(\d+)\s',                  # 問1
        r'【問(\d+)',                  # 【問1
        r'ID:\s*\d+\n.*?問(\d+)',      # ID: 86605\n...問1
        r'ユニットa\s+問(\d+)',         # ユニットa 問1
    ]

    found_numbers = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        for pattern in patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            for match in matches:
                try:
                    num = int(match)
                    if 1 <= num <= 100:  # 妥当な範囲
                        found_numbers.add(num)
                except:
                    pass

    # 49〜61の範囲をチェック
    missing_49_61 = []
    found_49_61 = []

    for n in range(49, 62):
        if n in found_numbers:
            found_49_61.append(n)
        else:
            missing_49_61.append(n)

    # 全体の問題番号範囲
    all_questions = sorted(found_numbers)
    min_q = min(all_questions) if all_questions else 0
    max_q = max(all_questions) if all_questions else 0

    return {
        'total': len(found_numbers),
        'range': f"{min_q}-{max_q}",
        'found_49_61': found_49_61,
        'missing_49_61': missing_49_61,
        'all_questions': all_questions,
    }


def analyze_r6_coverage(file_path):
    """
    R6年度の特別分析（gakkaBが完全欠落している理由）
    """
    doc = Document(file_path)

    # 学科A/Bの表記を検索
    gakka_a_count = 0
    gakka_b_count = 0

    full_text = '\n'.join([p.text for p in doc.paragraphs])

    gakka_a_count = len(re.findall(r'学科\s*A|問題A|gakkaA', full_text, re.IGNORECASE))
    gakka_b_count = len(re.findall(r'学科\s*B|問題B|gakkaB', full_text, re.IGNORECASE))

    return {
        'gakka_a_count': gakka_a_count,
        'gakka_b_count': gakka_b_count,
    }


def main():
    print("=== 原本ファイル 問49〜61 存在チェック ===\n")

    results = {}

    for year, filename in FILES.items():
        file_path = ORIGINAL_FOLDER / filename

        if not file_path.exists():
            print(f"❌ {year}: ファイルが見つかりません - {filename}")
            continue

        print(f"[{year}年度] {filename}")
        result = check_question_numbers(year, file_path)
        results[year] = result

        print(f"  総問題数: {result['total']}問")
        print(f"  問題番号範囲: {result['range']}")
        print(f"  問49〜61の存在: {len(result['found_49_61'])}/13問")

        if result['found_49_61']:
            print(f"    ✅ 存在: {result['found_49_61']}")

        if result['missing_49_61']:
            if len(result['missing_49_61']) == 13:
                print(f"    ❌ 問49〜61は**全て欠落**")
            else:
                print(f"    ❌ 欠落: {result['missing_49_61']}")

        # R6年度の特別分析
        if year == 'R6':
            r6_detail = analyze_r6_coverage(file_path)
            print(f"  R6年度 詳細:")
            print(f"    学科A言及: {r6_detail['gakka_a_count']}回")
            print(f"    学科B言及: {r6_detail['gakka_b_count']}回")
            if r6_detail['gakka_b_count'] == 0:
                print(f"    → **学科Bのデータなし**（学科Aのみ）")

        print()

    # サマリー
    print("=" * 80)
    print("【サマリー】")
    print("=" * 80)

    # 問49〜61の年度別存在状況
    print("\n問49〜61の存在状況:")
    for year in ['R1', 'R2', 'R3', 'R4', 'R5', 'R6', 'R7']:
        if year not in results:
            continue

        result = results[year]
        found_count = len(result['found_49_61'])

        if found_count == 0:
            status = "❌ 全欠落"
        elif found_count == 13:
            status = "✅ 全存在"
        else:
            status = f"⚠️  一部存在 ({found_count}/13)"

        print(f"  {year}: {status}")

    # 結論
    print("\n【結論】")

    all_missing = all(len(results[y]['found_49_61']) == 0 for y in results if y != 'R6')

    if all_missing:
        print("  🔴 問49〜61は**全年度で原本に存在しない**")
        print("     → 法規問題は解説が提供されていない可能性が高い")
        print("     → 過去問サイトでの収集可能性を確認する必要あり")
    else:
        print("  ⚠️  一部の年度では問49〜61が存在")
        print("     → 年度ごとに状況が異なる")

    # R6年度の結論
    if 'R6' in results:
        r6_result = results['R6']
        if r6_result['total'] < 100:
            print("\n  🔴 R6年度は**学科Aのみ**（学科Bのデータなし）")
            print("     → 35問（学科B全問）+ 一部学科A = 計70問欠落の説明")

    print("\n" + "=" * 80)


if __name__ == '__main__':
    main()
