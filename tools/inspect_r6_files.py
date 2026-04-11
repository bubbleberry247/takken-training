#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
R6年度の3つのWordファイルを詳細に調査

確認項目:
    - 実際の段落数、内容
    - 解説文の存在
    - 学科A/Bの区別
    - 問題番号の範囲
"""

import sys
import io
from pathlib import Path
from docx import Document
import re

# UTF-8出力設定
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

FILES = [
    r"C:\Users\masam\Downloads\１級土木-20260214T233348Z-1-001\１級土木\1級土木 解説文集 令和6年度.docx",
    r"C:\Users\masam\Downloads\１級土木-20260214T233348Z-1-001\１級土木\1級土木 解説文集 令和6年度(1).docx",
    r"C:\Users\masam\Downloads\１級土木-20260214T233348Z-1-001\１級土木\1級土木 解説文集 令和6年度(2).docx",
]

def analyze_file(file_path):
    """
    Wordファイルを詳細に分析
    """
    doc = Document(file_path)

    # 基本統計
    total_paras = len(doc.paragraphs)

    # 問題番号を検出
    problem_numbers = set()
    gakka_a_refs = 0
    gakka_b_refs = 0
    explanation_count = 0

    # 解説パターン
    explanation_patterns = [
        r'【解説】',
        r'【正解】',
        r'選択肢\d+\.',
        r'適当|不適当|誤り|正しい',
    ]

    # サンプルテキスト（最初の20段落）
    sample_text = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # サンプル収集（最初の20段落）
        if i < 20:
            sample_text.append(f"{i+1}: {text[:80]}...")

        # 問題番号検出
        problem_match = re.findall(r'問(\d+)[\s（]|【問(\d+)', text)
        for match in problem_match:
            num_str = match[0] or match[1]
            if num_str:
                try:
                    num = int(num_str)
                    if 1 <= num <= 100:
                        problem_numbers.add(num)
                except:
                    pass

        # 学科A/B検出
        if re.search(r'学科\s*A|問題A|gakkaA', text, re.IGNORECASE):
            gakka_a_refs += 1
        if re.search(r'学科\s*B|問題B|gakkaB', text, re.IGNORECASE):
            gakka_b_refs += 1

        # 解説パターン検出
        for pattern in explanation_patterns:
            if re.search(pattern, text):
                explanation_count += 1
                break

    # 問題番号範囲
    if problem_numbers:
        min_q = min(problem_numbers)
        max_q = max(problem_numbers)
        problem_range = f"{min_q}-{max_q}"
    else:
        problem_range = "検出なし"

    return {
        'file_name': Path(file_path).name,
        'file_size': Path(file_path).stat().st_size,
        'total_paras': total_paras,
        'problem_count': len(problem_numbers),
        'problem_range': problem_range,
        'problem_numbers': sorted(problem_numbers),
        'gakka_a_refs': gakka_a_refs,
        'gakka_b_refs': gakka_b_refs,
        'explanation_count': explanation_count,
        'sample_text': sample_text,
    }


def main():
    print("=== R6年度 Wordファイル詳細調査 ===\n")

    results = []

    for file_path in FILES:
        if not Path(file_path).exists():
            print(f"❌ ファイルが見つかりません: {file_path}")
            continue

        print(f"[調査中] {Path(file_path).name}")
        result = analyze_file(file_path)
        results.append(result)

        print(f"  ファイルサイズ: {result['file_size']:,} bytes ({result['file_size']/1024:.1f} KB)")
        print(f"  総段落数: {result['total_paras']}")
        print(f"  検出された問題数: {result['problem_count']}問")
        print(f"  問題番号範囲: {result['problem_range']}")
        print(f"  学科A言及: {result['gakka_a_refs']}回")
        print(f"  学科B言及: {result['gakka_b_refs']}回")
        print(f"  解説パターン検出: {result['explanation_count']}箇所")
        print()

    # 比較分析
    print("=" * 80)
    print("【ファイル比較】")
    print("=" * 80)

    if len(results) >= 2:
        # メインファイルと(1)の比較
        main = results[0]
        alt1 = results[1]

        print(f"\n{main['file_name']} vs {alt1['file_name']}:")
        print(f"  サイズ差: {abs(main['file_size'] - alt1['file_size'])} bytes")
        print(f"  段落数差: {abs(main['total_paras'] - alt1['total_paras'])}")
        print(f"  問題数差: {abs(main['problem_count'] - alt1['problem_count'])}")

        if main['file_size'] == alt1['file_size']:
            print(f"  → **完全に同一ファイルの可能性が高い**")
        elif abs(main['file_size'] - alt1['file_size']) < 1000:
            print(f"  → ほぼ同一（微差）")
        else:
            print(f"  → 内容が異なる")

    if len(results) >= 3:
        # (2)との比較
        alt2 = results[2]
        main = results[0]

        print(f"\n{main['file_name']} vs {alt2['file_name']}:")
        print(f"  サイズ差: {abs(main['file_size'] - alt2['file_size'])} bytes ({abs(main['file_size'] - alt2['file_size'])/main['file_size']*100:.1f}%)")
        print(f"  段落数差: {abs(main['total_paras'] - alt2['total_paras'])}")
        print(f"  問題数差: {abs(main['problem_count'] - alt2['problem_count'])}")

        if alt2['file_size'] < main['file_size'] * 0.7:
            print(f"  → **(2)は不完全データの可能性が高い**（{alt2['file_size']/main['file_size']*100:.1f}%のサイズ）")

    # 詳細内容表示
    print("\n" + "=" * 80)
    print("【詳細内容サンプル（最初の20段落）】")
    print("=" * 80)

    for result in results:
        print(f"\n## {result['file_name']}")
        for line in result['sample_text']:
            print(f"  {line}")

    # 問題番号リスト
    print("\n" + "=" * 80)
    print("【検出された問題番号】")
    print("=" * 80)

    for result in results:
        print(f"\n{result['file_name']}:")
        if result['problem_numbers']:
            # 連続範囲を見やすく表示
            nums = result['problem_numbers']
            print(f"  総数: {len(nums)}問")
            print(f"  範囲: {min(nums)}-{max(nums)}")

            # 欠落している番号を検出
            full_range = set(range(min(nums), max(nums) + 1))
            missing = sorted(full_range - set(nums))

            if missing:
                print(f"  欠落: {missing}")
            else:
                print(f"  欠落: なし（連続）")

            # 問49〜61の存在チェック
            range_49_61 = set(range(49, 62))
            found_49_61 = sorted(range_49_61 & set(nums))
            missing_49_61 = sorted(range_49_61 - set(nums))

            print(f"  問49〜61: {len(found_49_61)}/13問")
            if found_49_61:
                print(f"    存在: {found_49_61}")
            if missing_49_61:
                print(f"    欠落: {missing_49_61}")
        else:
            print(f"  問題番号検出なし")


if __name__ == '__main__':
    main()
