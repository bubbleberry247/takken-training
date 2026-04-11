#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
説明文完全欠落97問の原本詳細調査

対象:
    - R1gakkaA-049〜061（13問）
    - R2gakkaA-049〜061（13問）
    - R3gakkaA-049〜061（13問）
    - R4gakkaA-049〜061（13問）
    - R5gakkaA-049〜061（13問）
    - R6gakkaA-036〜066（31問）
    - R7gakkaA-005（1問）
"""

import sys
import io
from pathlib import Path
from docx import Document
import re

# UTF-8出力設定
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

ORIGINAL_FOLDER = Path(r"C:\Users\masam\Downloads\１級土木-20260214T233348Z-1-001\１級土木")

# 推奨ファイル
FILES = {
    'R1': "1級土木 解説文集 令和元年度.docx",
    'R2': "1級土木 解説文集 令和2年度.docx",
    'R3': "1級土木 解説文集 令和3年度(1).docx",
    'R4': "1級土木 解説文集 令和4年度.docx",
    'R5': "1級土木 解説文集 令和5年度.docx",
    'R6': "1級土木 解説文集 令和6年度.docx",
    'R7': "1級土木 解説文集 令和7年度.docx",
}

# 調査対象問題
TARGET_PROBLEMS = {
    'R1': list(range(49, 62)),  # 049-061
    'R2': list(range(49, 62)),
    'R3': list(range(49, 62)),
    'R4': list(range(49, 62)),
    'R5': list(range(49, 62)),
    'R6': list(range(36, 67)),  # 036-066
    'R7': [5],                   # 005のみ
}


def search_problem_in_file(year, file_path, target_numbers):
    """
    Wordファイルから特定の問題番号を検索

    Returns:
        dict: {問題番号: {found: bool, context: str}}
    """
    doc = Document(file_path)
    full_text = '\n'.join([p.text for p in doc.paragraphs])

    results = {}

    for num in target_numbers:
        # 問題番号パターン（複数形式）
        patterns = [
            rf'問{num}[\s（]',
            rf'【問{num}',
            rf'問題{num}',
            rf'No\.{num}[\s）]',
        ]

        found = False
        context = ""

        for pattern in patterns:
            matches = list(re.finditer(pattern, full_text))
            if matches:
                found = True
                # コンテキスト取得（マッチ前後100文字）
                match = matches[0]
                start = max(0, match.start() - 100)
                end = min(len(full_text), match.end() + 200)
                context = full_text[start:end].replace('\n', ' ')
                break

        # 解説パターンの検索
        has_explanation = False
        if found:
            # 問題番号の後に解説パターンがあるか
            explanation_patterns = [
                r'【解説】',
                r'選択肢\d+\.',
                r'適当|不適当|誤り|正しい',
            ]

            # 問題番号から次の問題番号までのテキストを抽出
            problem_section = ""
            for i, para in enumerate(doc.paragraphs):
                if f'問{num}' in para.text or f'【問{num}' in para.text:
                    # 次の問題まで収集
                    section_paras = []
                    for j in range(i, min(i+50, len(doc.paragraphs))):
                        text = doc.paragraphs[j].text.strip()
                        section_paras.append(text)
                        # 次の問題検出（区切り線または次の問題番号）
                        if j > i and ('────' in text or re.search(rf'問{num+1}[\s（]', text)):
                            break
                    problem_section = '\n'.join(section_paras)
                    break

            # 解説パターンをチェック
            for exp_pattern in explanation_patterns:
                if re.search(exp_pattern, problem_section):
                    has_explanation = True
                    break

        results[num] = {
            'found': found,
            'has_explanation': has_explanation,
            'context': context[:200],
            'qid': f"{year}gakkaA-{num:03d}",
        }

    return results


def main():
    print("=== 説明文完全欠落97問の原本詳細調査 ===\n")

    all_results = {}

    for year, filename in FILES.items():
        if year not in TARGET_PROBLEMS:
            continue

        file_path = ORIGINAL_FOLDER / filename
        if not file_path.exists():
            print(f"❌ {year}: ファイルが見つかりません")
            continue

        print(f"[{year}年度] {filename}")
        target_nums = TARGET_PROBLEMS[year]
        results = search_problem_in_file(year, file_path, target_nums)
        all_results[year] = results

        # 結果サマリー
        found_count = sum(1 for r in results.values() if r['found'])
        has_exp_count = sum(1 for r in results.values() if r['has_explanation'])

        print(f"  対象問題: {len(target_nums)}問")
        print(f"  検出: {found_count}/{len(target_nums)}問")
        print(f"  解説あり: {has_exp_count}/{len(target_nums)}問")

        # 詳細（見つからなかった問題のみ）
        not_found = [num for num, r in results.items() if not r['found']]
        if not_found:
            print(f"  ❌ 未検出: {not_found}")

        no_exp = [num for num, r in results.items() if r['found'] and not r['has_explanation']]
        if no_exp:
            print(f"  ⚠️  解説なし: {no_exp}")

        print()

    # 総合レポート
    print("=" * 80)
    print("【総合レポート】")
    print("=" * 80)

    total_target = sum(len(TARGET_PROBLEMS[y]) for y in all_results)
    total_found = sum(sum(1 for r in results.values() if r['found']) for results in all_results.values())
    total_has_exp = sum(sum(1 for r in results.values() if r['has_explanation']) for results in all_results.values())

    print(f"\n対象問題: {total_target}問")
    print(f"原本で検出: {total_found}問 ({total_found/total_target*100:.1f}%)")
    print(f"解説あり: {total_has_exp}問 ({total_has_exp/total_target*100:.1f}%)")
    print(f"解説なし: {total_found - total_has_exp}問")

    # 年度別サマリー
    print("\n【年度別サマリー】")
    print("| 年度 | 対象 | 検出 | 解説あり | 解説なし | 状態 |")
    print("|------|------|------|---------|---------|------|")

    for year in sorted(all_results.keys()):
        results = all_results[year]
        target = len(results)
        found = sum(1 for r in results.values() if r['found'])
        has_exp = sum(1 for r in results.values() if r['has_explanation'])
        no_exp = found - has_exp

        if has_exp == target:
            status = "✅ 全問解説あり"
        elif has_exp > 0:
            status = f"⚠️  一部解説あり ({has_exp}/{target})"
        else:
            status = "❌ 解説なし"

        print(f"| {year} | {target}問 | {found}問 | {has_exp}問 | {no_exp}問 | {status} |")

    # 結論
    print("\n【結論】")

    if total_has_exp == total_target:
        print("✅ 全97問の解説が原本に存在します！")
        print("   → パーサー改善で全て回収可能")
    elif total_has_exp > 0:
        print(f"⚠️  {total_has_exp}問の解説が原本に存在")
        print(f"   → パーサー改善で{total_has_exp}問回収可能")
        print(f"   → 残り{total_target - total_has_exp}問は別の方法が必要")
    else:
        print("❌ 原本に解説が存在しない可能性")
        print("   → 過去問サイトからの収集を検討")

    print("=" * 80)

    # 詳細レポートをファイルに保存
    output_file = Path(__file__).parent / "missing_97_investigation.json"
    import json
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_results, f, ensure_ascii=False, indent=2)

    print(f"\n✓ 詳細レポート保存: {output_file.name}")

    return all_results


if __name__ == '__main__':
    main()
