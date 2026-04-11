#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Wordドキュメントの構造を調査
"""

import sys
import io
from docx import Document
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DOCX_PATH = Path(r"C:\Users\masam\Downloads\不足説明文.docx")

doc = Document(DOCX_PATH)

print(f'📄 ドキュメント: {DOCX_PATH.name}')
print('='*80)
print(f'総段落数: {len(doc.paragraphs)}')
print(f'総テーブル数: {len(doc.tables)}')
print()

# 段落のサンプル
print('段落サンプル（最初の20個）:')
for i, para in enumerate(doc.paragraphs[:20], 1):
    text = para.text.strip()
    if text:
        print(f'{i}. {text[:100]}')

# テーブルの構造
if doc.tables:
    print(f'\n\n📊 テーブル情報:')
    for table_idx, table in enumerate(doc.tables, 1):
        print(f'\nテーブル {table_idx}:')
        print(f'  行数: {len(table.rows)}')
        print(f'  列数: {len(table.columns) if table.rows else 0}')

        # 最初の3行を表示
        if len(table.rows) > 0:
            print(f'\n  最初の3行:')
            for row_idx, row in enumerate(table.rows[:3], 1):
                cells = [cell.text.strip()[:30] for cell in row.cells]
                print(f'    行{row_idx}: {" | ".join(cells)}')
