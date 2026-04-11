#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
元のドキュメント（無題のドキュメント.docx）を詳細調査
"""

import sys
import io
from docx import Document
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DOCX_PATH = Path(r"C:\Users\masam\Downloads\無題のドキュメント.docx")

doc = Document(DOCX_PATH)

print(f'📄 ドキュメント: {DOCX_PATH.name}')
print('='*80)
print(f'総段落数: {len(doc.paragraphs)}')
print(f'総テーブル数: {len(doc.tables)}')
print()

# テーブルの内容を詳細表示
if doc.tables:
    for table_idx, table in enumerate(doc.tables, 1):
        print(f'\nテーブル {table_idx}:')
        print(f'  行数: {len(table.rows)}')
        print(f'  列数: {len(table.columns) if table.rows else 0}')

        # 全行を表示（最大20行）
        if len(table.rows) > 0:
            print(f'\n  全行表示（最大20行）:')
            for row_idx, row in enumerate(table.rows[:20]):
                cells = [cell.text.strip() for cell in row.cells]
                # qId列（最初の列）を確認
                qid = cells[0] if cells else ''
                status = cells[7] if len(cells) > 7 else ''
                print(f'    行{row_idx+1}: qId={qid[:20]}, 状態={status}')

# 段落にMarkdown形式のテーブルがあるか確認
print(f'\n\n📝 Markdown形式のテーブルを探索:')
table_lines = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith('|') and 'qId' in text or 'R6gakkaA' in text or 'R5gakkaA' in text or 'R7gakkaA' in text:
        table_lines.append(text[:100])

if table_lines:
    print(f'  見つかりました！（最初の10行）:')
    for i, line in enumerate(table_lines[:10], 1):
        print(f'    {i}. {line}')
else:
    print(f'  見つかりませんでした。')
