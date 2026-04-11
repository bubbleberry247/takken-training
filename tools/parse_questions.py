"""Parse docx question files + answer docx into structured JSON.

Input directory layout (example):
  docx_out_v2/
    R7gakkaA_mondai.docx
    R7gakkaB_mondai.docx
    ...
    解答/
      R7gakka_kaitou.docx
      ...

Usage:
  python parse_questions.py --input-dir "C:/.../docx_out_v2" --output questions_parsed.json

Notes:
- Questions are detected by lines like: 【No. 1】 ...
- Choices are detected by lines like: (1) ... (2) ...
- This script keeps formatting minimal (paragraphs are joined by spaces).
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document


QUESTION_START_RE = re.compile(r"^【No\.\s*(\d+)】\s*(.*)$")
CHOICE_START_RE = re.compile(r"^\((\d+)\)\s*(.*)$")
ANSWER_DIGITS_RE = re.compile(r"解答\s*([0-9]+)")

try:
    # Optional fallback (R4 answer docx is currently image-only).
    from integrate_manual_answers import R4_ANSWERS, R5_ANSWERS, R6_ANSWERS  # type: ignore

    MANUAL_ANSWERS = {"R4": R4_ANSWERS, "R5": R5_ANSWERS, "R6": R6_ANSWERS}
except Exception:
    MANUAL_ANSWERS = {}


def _join_parts(parts: List[str]) -> str:
    # Join docx paragraphs into a single string while keeping it import-friendly.
    return " ".join(p.strip() for p in parts if p and p.strip()).strip()


def _is_math_like(text: str) -> bool:
    # Heuristic: short, mostly non-CJK (often numerator like "4P").
    t = text.strip()
    if not t:
        return False
    if len(t) > 16:
        return False
    # If it contains any CJK, treat as not math-like.
    if re.search(r"[\u3040-\u30ff\u3400-\u9fff]", t):
        return False
    return True


def _finalize_question(q_num: int, body_lines: List[str], line_start: int) -> Dict:
    stem_parts: List[str] = []
    choices: Dict[int, List[str]] = {}

    current_choice: int | None = None

    for line in body_lines:
        m = CHOICE_START_RE.match(line)
        if m:
            try:
                cnum = int(m.group(1))
            except ValueError:
                cnum = -1
            if 1 <= cnum <= 5:
                # Move last math-like stem token into the first choice.
                if current_choice is None and stem_parts and _is_math_like(stem_parts[-1]):
                    moved = stem_parts.pop()
                    choices.setdefault(cnum, []).append(moved)

                current_choice = cnum
                rest = m.group(2).strip()
                if rest:
                    choices.setdefault(cnum, []).append(rest)
                else:
                    choices.setdefault(cnum, [])
                continue

        if current_choice is None:
            stem_parts.append(line)
        else:
            choices.setdefault(current_choice, []).append(line)

    choice_texts: Dict[int, str] = {k: _join_parts(v) for k, v in choices.items()}

    return {
        "qNum": q_num,
        "stem": _join_parts(stem_parts),
        "choices": choice_texts,
        "lineStart": line_start,
    }


def parse_question_file(docx_path: Path) -> List[Dict]:
    doc = Document(str(docx_path))
    questions: List[Dict] = []

    current_num: int | None = None
    current_lines: List[str] = []
    current_start: int = 0

    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue

        m = QUESTION_START_RE.match(text)
        if m:
            if current_num is not None:
                questions.append(_finalize_question(current_num, current_lines, current_start))

            current_num = int(m.group(1))
            current_start = i
            current_lines = []

            rest = m.group(2).strip()
            if rest:
                current_lines.append(rest)
            continue

        if current_num is None:
            # Skip cover/instructions before first question.
            continue

        current_lines.append(text)

    if current_num is not None:
        questions.append(_finalize_question(current_num, current_lines, current_start))

    return questions


def extract_answers(docx_path: Path) -> List[str]:
    doc = Document(str(docx_path))
    groups: List[str] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        m = ANSWER_DIGITS_RE.search(text)
        if m:
            groups.append(m.group(1))
            continue

        # Some years omit the "解答" label; capture long digit runs of [1-5] only.
        if "問題番号" in text or "正答肢" in text:
            continue

        compact = re.sub(r"\s+", "", text)
        m0 = re.search(r"[0-9]", compact)
        if not m0:
            continue

        digits = compact[m0.start() :]
        if re.fullmatch(r"[1-5]{10,}", digits):
            groups.append(digits)

    return groups


def build_answer_map(groups: List[str], section: str, expected_count: int) -> Dict[int, int]:
    if len(groups) < 5:
        raise ValueError(f"Answer docx seems invalid: expected >= 5 digit groups, got {len(groups)}")

    if section.upper() == "A":
        digits = "".join(groups[:3])
    elif section.upper() == "B":
        digits = "".join(groups[3:5])
    else:
        raise ValueError(f"Unknown section: {section}")

    if len(digits) < expected_count:
        raise ValueError(
            f"Not enough answer digits for section {section}: expected {expected_count}, got {len(digits)}"
        )

    out: Dict[int, int] = {}
    for i in range(expected_count):
        out[i + 1] = int(digits[i])
    return out


def main() -> None:
    parser = argparse.ArgumentParser(description="Parse question + answer docx files")
    parser.add_argument("--input-dir", required=True, help="Directory containing *_mondai.docx and 解答/*.docx")
    parser.add_argument("--output", default="questions_parsed.json", help="Output JSON file")
    parser.add_argument("--test", action="store_true", help="Test mode: only parse R7gakkaA")
    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    answer_dir = input_dir / "解答"

    years = ["R1", "R2", "R3", "R4", "R5", "R6", "R7"]
    sections = ["A", "B"]

    if args.test:
        years = ["R7"]
        sections = ["A"]

    all_data: Dict[str, List[Dict]] = {}

    for year in years:
        answer_file = answer_dir / f"{year}gakka_kaitou.docx"
        if not answer_file.exists():
            print(f"Warning: answer file not found: {answer_file}")
            continue

        groups = extract_answers(answer_file)
        manual = MANUAL_ANSWERS.get(year)

        for section in sections:
            q_file = input_dir / f"{year}gakka{section}_mondai.docx"
            if not q_file.exists():
                print(f"Warning: question file not found: {q_file}")
                continue

            questions = parse_question_file(q_file)
            if not questions:
                print(f"Warning: no questions parsed: {q_file}")
                continue

            max_num = max(q["qNum"] for q in questions)
            ans_map: Dict[int, int] = {}
            if len(groups) >= 5:
                ans_map = build_answer_map(groups, section, max_num)
            elif manual and section in manual:
                seq = manual[section]
                if len(seq) < max_num:
                    raise ValueError(
                        f"Manual answers too short for {year}gakka{section}: expected {max_num}, got {len(seq)}"
                    )
                ans_map = {i + 1: int(seq[i]) for i in range(max_num)}
                print(f"Info: using manual answers for {year}gakka{section} ({max_num})")
            else:
                print(f"Warning: no answers available for {year}gakka{section} (parsed=0, manual=none)")

            for q in questions:
                q_num = int(q["qNum"])
                if q_num in ans_map:
                    q["correct"] = ans_map[q_num]

            key = f"{year}gakka{section}"
            all_data[key] = questions

            print(f"{key}: questions={len(questions)} maxNo={max_num} answers={len(ans_map)}")

    out_path = Path(args.output)
    with out_path.open("w", encoding="utf-8") as f:
        json.dump(all_data, f, ensure_ascii=False, indent=2)

    total = sum(len(v) for v in all_data.values())
    print(f"Saved: {out_path} (total questions: {total})")


if __name__ == "__main__":
    main()
